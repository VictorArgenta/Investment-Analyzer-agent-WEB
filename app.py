"""
Investment Analyzer Web — Flask + Multi-model agent (Claude / Gemini)
Análisis institucional de inversión vía agente con búsqueda web.
"""
import io
import json
import os
import queue
import re
import threading
import uuid
from datetime import datetime
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen

from anthropic import Anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from flask import Flask, Response, jsonify, render_template, request, send_file

load_dotenv()

app = Flask(__name__)

# ---------------------------------------------------------------------------
# Catálogo de modelos disponibles
# ---------------------------------------------------------------------------
# provider: "anthropic" (de pago) | "google" (gratuito)
MODEL_CATALOG: dict[str, dict] = {
    # Anthropic Claude — de pago
    "claude-opus-4-6": {
        "provider": "anthropic",
        "label": "Claude Opus 4.6",
        "tier": "Máxima calidad",
        "paid": True,
    },
    "claude-sonnet-4-6": {
        "provider": "anthropic",
        "label": "Claude Sonnet 4.6",
        "tier": "Recomendado",
        "paid": True,
    },
    "claude-haiku-4-5-20251001": {
        "provider": "anthropic",
        "label": "Claude Haiku 4.5",
        "tier": "Más económico",
        "paid": True,
    },
    # Google Gemini — gratuito
    "gemini-2.5-flash": {
        "provider": "google",
        "label": "Gemini 2.5 Flash",
        "tier": "Recomendado gratuito",
        "paid": False,
    },
    "gemini-2.5-flash-lite": {
        "provider": "google",
        "label": "Gemini 2.5 Flash Lite",
        "tier": "Ligero gratuito",
        "paid": False,
    },
}
DEFAULT_MODEL = "claude-sonnet-4-6"

# Herramienta web_search para Anthropic
ANTHROPIC_WEB_SEARCH_TOOL = {
    "type": "web_search_20250305",
    "name": "web_search",
    "max_uses": 15,
}
MAX_TOKENS = 16000
MAX_ITERATIONS = 25

# Ruta del system prompt
PROMPT_PATH = Path(__file__).parent / "prompts" / "investment_agent.md"

# Estado por sesión — cola para SSE + informe final para descarga
sessions: dict[str, queue.Queue] = {}
reports: dict[str, dict] = {}
state_lock = threading.Lock()


def load_system_prompt() -> str:
    """Lee el system prompt desde el fichero en prompts/."""
    with open(PROMPT_PATH, "r", encoding="utf-8") as f:
        return f.read()


def emit(q: queue.Queue, event_type: str, **payload) -> None:
    """Envía un evento a la cola SSE de la sesión."""
    q.put({"type": event_type, **payload})


def build_user_message(ticker: str, search_tool_label: str) -> str:
    return (
        f"Realiza un Equity Research institucional completo sobre la empresa cotizada "
        f"con ticker {ticker}. Sigue rigurosamente los 8 pasos del proceso de investigación "
        f"usando {search_tool_label} para obtener datos actualizados y fuentes primarias "
        f"(SEC filings, Investor Relations, earnings calls). "
        f"Cuando hayas completado TODOS los pasos, redacta el informe final "
        f"con el formato institucional especificado."
    )


def store_report(session_id: str, ticker: str, report_text: str, model_name: str) -> None:
    with state_lock:
        reports[session_id] = {
            "ticker": ticker,
            "report": report_text,
            "model": model_name,
            "generated_at": datetime.utcnow().isoformat() + "Z",
        }


# ---------------------------------------------------------------------------
# Backend 1 — Anthropic Claude con web_search_20250305
# ---------------------------------------------------------------------------


def run_agent_anthropic(ticker: str, model_name: str, session_id: str, q: queue.Queue) -> None:
    """Bucle agentico completo sobre la API de Anthropic."""
    try:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            emit(q, "error", message="Falta ANTHROPIC_API_KEY en el entorno")
            return

        client = Anthropic(api_key=api_key)
        system_prompt = load_system_prompt()

        emit(q, "step", message=f"Iniciando análisis de {ticker} con {model_name}")
        emit(q, "step", message="Cargando system prompt de Equity Research")

        messages = [{"role": "user", "content": build_user_message(ticker, "web_search")}]

        iteration = 0
        final_response = None

        while iteration < MAX_ITERATIONS:
            iteration += 1
            emit(q, "step", message=f"Llamada al modelo #{iteration}")

            response = client.messages.create(
                model=model_name,
                max_tokens=MAX_TOKENS,
                system=system_prompt,
                tools=[ANTHROPIC_WEB_SEARCH_TOOL],
                messages=messages,
            )
            final_response = response

            for block in response.content:
                btype = getattr(block, "type", None)

                if btype == "text":
                    text = getattr(block, "text", "") or ""
                    if text.strip():
                        preview = text.strip()
                        if len(preview) > 280:
                            preview = preview[:280] + "…"
                        emit(q, "thinking", message=preview)

                elif btype == "server_tool_use":
                    tool_name = getattr(block, "name", "tool")
                    tool_input = getattr(block, "input", {}) or {}
                    query = tool_input.get("query") if isinstance(tool_input, dict) else None
                    if query:
                        emit(q, "tool_use", message=f"web_search · {query}")
                    else:
                        emit(q, "tool_use", message=f"Invocando {tool_name}")

                elif btype == "web_search_tool_result":
                    content = getattr(block, "content", None)
                    count = len(content) if isinstance(content, list) else 0
                    emit(q, "tool_result", message=f"Resultados recibidos ({count} fuentes)")

                elif btype == "tool_use":
                    emit(q, "tool_use", message=f"Herramienta local: {getattr(block, 'name', '')}")

            messages.append({"role": "assistant", "content": response.content})

            stop_reason = response.stop_reason
            emit(q, "step", message=f"stop_reason = {stop_reason}")

            if stop_reason == "tool_use":
                tool_results = []
                for block in response.content:
                    if getattr(block, "type", None) == "tool_use":
                        tool_results.append({
                            "type": "tool_result",
                            "tool_use_id": block.id,
                            "content": "Herramienta no implementada en este entorno.",
                            "is_error": True,
                        })
                if not tool_results:
                    break
                messages.append({"role": "user", "content": tool_results})
                continue

            if stop_reason == "pause_turn":
                emit(q, "step", message="Pausa de turno — continuando")
                continue

            break

        report_text = ""
        if final_response is not None:
            report_text = "\n\n".join(
                getattr(b, "text", "") for b in final_response.content
                if getattr(b, "type", None) == "text"
            ).strip()

        if not report_text:
            report_text = "El agente terminó sin producir un informe de texto."

        store_report(session_id, ticker, report_text, model_name)
        emit(q, "done", report=report_text, iterations=iteration, session_id=session_id)

    except Exception as exc:  # noqa: BLE001
        emit(q, "error", message=f"{type(exc).__name__}: {exc}")


# ---------------------------------------------------------------------------
# Backend 2 — Google Gemini con Google Search grounding
# ---------------------------------------------------------------------------


GEMINI_API_BASE = "https://generativelanguage.googleapis.com/v1beta/models"


def _gemini_generate(model_name: str, api_key: str, system_prompt: str, user_message: str,
                     tool_results: list | None = None) -> dict:
    """
    Llamada REST a la Gemini API con el tool `google_search` (grounding nativo de 2.5).
    Usamos REST porque el wrapper google-generativeai 0.8.x aún no expone
    el campo `google_search` del proto (sólo el legacy `google_search_retrieval`).
    """
    url = f"{GEMINI_API_BASE}/{model_name}:generateContent?key={api_key}"
    body = {
        "systemInstruction": {"parts": [{"text": system_prompt}]},
        "contents": [{"role": "user", "parts": [{"text": user_message}]}],
        "tools": [{"google_search": {}}],
        "generationConfig": {
            "temperature": 0.3,
            "maxOutputTokens": 8192,
        },
    }
    req = Request(
        url,
        data=json.dumps(body).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urlopen(req, timeout=300) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except HTTPError as exc:
        err_body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Gemini HTTP {exc.code}: {err_body[:500]}") from exc
    except URLError as exc:
        raise RuntimeError(f"Gemini network error: {exc.reason}") from exc


def run_agent_gemini(ticker: str, model_name: str, session_id: str, q: queue.Queue) -> None:
    """
    Ejecuta el análisis con Gemini 2.5 y Google Search grounding (tool `google_search`).
    """
    try:
        api_key = os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            emit(q, "error", message="Falta GOOGLE_API_KEY en el entorno")
            return

        system_prompt = load_system_prompt()

        emit(q, "step", message=f"Iniciando análisis de {ticker} con {model_name}")
        emit(q, "step", message="Cargando system prompt de Equity Research")

        user_message = build_user_message(ticker, "Google Search")

        emit(q, "step", message="Llamada al modelo Gemini con Google Search grounding")

        data = _gemini_generate(model_name, api_key, system_prompt, user_message)

        # Extraer texto y metadata de grounding del JSON de respuesta
        report_chunks: list[str] = []
        total_queries = 0
        total_sources = 0

        for candidate in data.get("candidates", []) or []:
            grounding = candidate.get("groundingMetadata") or {}
            for query in grounding.get("webSearchQueries", []) or []:
                total_queries += 1
                emit(q, "tool_use", message=f"google_search · {query}")

            chunks = grounding.get("groundingChunks", []) or []
            if chunks:
                total_sources += len(chunks)
                emit(q, "tool_result", message=f"Resultados recibidos ({len(chunks)} fuentes)")

            for part in (candidate.get("content") or {}).get("parts", []) or []:
                text = part.get("text")
                if text:
                    report_chunks.append(text)

            finish_reason = candidate.get("finishReason")
            if finish_reason:
                emit(q, "step", message=f"finishReason = {finish_reason}")

        report_text = "".join(report_chunks).strip()

        if not report_text:
            # En caso de prompt bloqueado o sin candidatos, dar pista
            block_reason = (data.get("promptFeedback") or {}).get("blockReason")
            if block_reason:
                report_text = f"La solicitud fue bloqueada por Gemini: {block_reason}"
            else:
                report_text = "El agente terminó sin producir un informe de texto."

        emit(q, "step", message=f"Grounding total: {total_queries} búsquedas, {total_sources} fuentes")

        store_report(session_id, ticker, report_text, model_name)
        emit(q, "done", report=report_text, iterations=1, session_id=session_id)

    except Exception as exc:  # noqa: BLE001
        emit(q, "error", message=f"{type(exc).__name__}: {exc}")


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------


def run_agent(ticker: str, model_name: str, session_id: str, q: queue.Queue) -> None:
    """Elige el backend según el proveedor del modelo."""
    info = MODEL_CATALOG.get(model_name)
    if info is None:
        emit(q, "error", message=f"Modelo no soportado: {model_name}")
        return

    provider = info["provider"]
    if provider == "anthropic":
        run_agent_anthropic(ticker, model_name, session_id, q)
    elif provider == "google":
        run_agent_gemini(ticker, model_name, session_id, q)
    else:
        emit(q, "error", message=f"Proveedor desconocido: {provider}")


# ---------------------------------------------------------------------------
# Construcción del documento Word
# ---------------------------------------------------------------------------

_MD_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _add_rich_paragraph(doc: Document, text: str, *, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    cursor = 0
    for match in _MD_BOLD.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _strip_md(text: str) -> str:
    text = _MD_BOLD.sub(r"\1", text)
    text = text.replace("`", "")
    return text


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        raw = lines[i].strip()
        if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", raw):
            i += 1
            continue
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def build_docx(ticker: str, report_text: str, generated_at: str, model_name: str = "") -> io.BytesIO:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"Investment Research Report — {ticker}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x5F)

    meta = doc.add_paragraph()
    meta_text = f"Generado: {generated_at}  ·  Equity Research AI"
    if model_name:
        meta_text += f"  ·  Modelo: {model_name}"
    meta_run = meta.add_run(meta_text)
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

    doc.add_paragraph()

    lines = report_text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if not line.strip():
            i += 1
            continue

        stripped = line.lstrip()

        if stripped.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            rows, next_i = _parse_table(lines, i)
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(ncols):
                        cell_text = _strip_md(row[c_idx]) if c_idx < len(row) else ""
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        run = para.add_run(cell_text)
                        if r_idx == 0:
                            run.bold = True
                doc.add_paragraph()
            i = next_i
            continue

        if stripped.startswith("#### "):
            doc.add_heading(_strip_md(stripped[5:]), level=4)
        elif stripped.startswith("### "):
            doc.add_heading(_strip_md(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(_strip_md(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(_strip_md(stripped[2:]), level=1)
        elif stripped.startswith(("- ", "* ", "• ")):
            _add_rich_paragraph(doc, stripped[2:].lstrip(), style="List Bullet")
        elif re.match(r"^\d+[.\)]\s+", stripped):
            content = re.sub(r"^\d+[.\)]\s+", "", stripped)
            _add_rich_paragraph(doc, content, style="List Number")
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 40)
        else:
            _add_rich_paragraph(doc, stripped)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------------


@app.route("/")
def index():
    """Formulario con ticker y selector de modelo."""
    # Pasamos el catálogo agrupado a la plantilla
    grouped = {"paid": [], "free": []}
    for model_id, info in MODEL_CATALOG.items():
        entry = {"id": model_id, **info}
        grouped["paid" if info["paid"] else "free"].append(entry)
    return render_template("index.html", models=grouped, default_model=DEFAULT_MODEL)


@app.route("/analyze", methods=["POST"])
def analyze():
    """Lanza el agente con el modelo elegido y devuelve el session_id."""
    data = request.get_json(silent=True) or request.form
    ticker = (data.get("ticker") or "").strip().upper()
    model_name = (data.get("model") or DEFAULT_MODEL).strip()

    if not ticker:
        return jsonify({"error": "El ticker es obligatorio"}), 400
    if model_name not in MODEL_CATALOG:
        return jsonify({"error": f"Modelo no soportado: {model_name}"}), 400

    session_id = str(uuid.uuid4())
    q: queue.Queue = queue.Queue()
    with state_lock:
        sessions[session_id] = q

    thread = threading.Thread(
        target=run_agent,
        args=(ticker, model_name, session_id, q),
        daemon=True,
    )
    thread.start()

    return jsonify({
        "session_id": session_id,
        "ticker": ticker,
        "model": model_name,
        "provider": MODEL_CATALOG[model_name]["provider"],
    })


@app.route("/stream/<session_id>")
def stream(session_id: str):
    """Stream SSE con los eventos del agente."""

    def generate():
        with state_lock:
            q = sessions.get(session_id)
        if q is None:
            yield f"data: {json.dumps({'type': 'error', 'message': 'Sesión no encontrada'})}\n\n"
            return

        try:
            while True:
                try:
                    event = q.get(timeout=300)
                except queue.Empty:
                    yield ": keep-alive\n\n"
                    continue

                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

                if event.get("type") in ("done", "error"):
                    break
        finally:
            with state_lock:
                sessions.pop(session_id, None)

    headers = {
        "Content-Type": "text/event-stream",
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
        "Connection": "keep-alive",
    }
    return Response(generate(), headers=headers)


@app.route("/download/<session_id>")
def download(session_id: str):
    """Devuelve el informe como documento Word (.docx)."""
    with state_lock:
        data = reports.get(session_id)
    if not data:
        return jsonify({"error": "Informe no disponible o sesión expirada"}), 404

    buf = build_docx(data["ticker"], data["report"], data["generated_at"], data.get("model", ""))
    stamp = datetime.utcnow().strftime("%Y%m%d_%H%M")
    filename = f"research_{data['ticker']}_{stamp}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=False, threaded=True)
